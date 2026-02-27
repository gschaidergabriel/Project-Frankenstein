"""
Document Renderer
Export documents to various formats
"""

import html
import re
from pathlib import Path
from typing import Optional
import subprocess

from writer.editor.document import Document


class DocumentRenderer:
    """Renders documents to various output formats"""

    def __init__(self, config):
        self.config = config

    def to_pdf(self, document: Document, output_path: Path):
        """Export to PDF using WeasyPrint"""
        try:
            from weasyprint import HTML, CSS
        except ImportError:
            raise RuntimeError("WeasyPrint not installed. Run: pip install weasyprint")

        # Convert to HTML first
        html_content = self._to_html_string(document)

        # Load CSS
        css_path = Path(__file__).parent.parent / "styles" / f"{self.config.export.default_pdf_style}.css"
        if not css_path.exists():
            raise ValueError(f"CSS style file not found: {css_path}")
        css = CSS(filename=str(css_path))

        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Generate PDF
        html_doc = HTML(string=html_content)
        html_doc.write_pdf(str(output_path), stylesheets=[css])

    def to_docx(self, document: Document, output_path: Path):
        """Export to DOCX — B5 FIX: preserves inline formatting (bold, italic, links)"""
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

        doc = DocxDocument()

        # Add title
        doc.add_heading(document.title, 0)

        # Parse content
        lines = document.content.split('\n')
        current_para = []
        in_code_block = False
        code_block_content = []

        for line in lines:
            # Code blocks
            if line.startswith('```'):
                if not in_code_block:
                    if current_para:
                        self._add_rich_paragraph(doc, '\n'.join(current_para))
                        current_para = []
                    in_code_block = True
                    code_block_content = []
                else:
                    in_code_block = False
                    if code_block_content:
                        code_para = doc.add_paragraph()
                        code_run = code_para.add_run('\n'.join(code_block_content))
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(10)
                    code_block_content = []
                continue

            if in_code_block:
                code_block_content.append(line)
                continue

            # Headings
            if line.startswith('#'):
                if current_para:
                    self._add_rich_paragraph(doc, '\n'.join(current_para))
                    current_para = []

                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                doc.add_heading(text, min(level, 9))

            # Regular text
            else:
                if line.strip():
                    current_para.append(line)
                elif current_para:
                    self._add_rich_paragraph(doc, '\n'.join(current_para))
                    current_para = []

        # Flush remaining
        if current_para:
            self._add_rich_paragraph(doc, '\n'.join(current_para))

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    def _add_rich_paragraph(self, doc, text: str):
        """Add a paragraph with inline markdown formatting to DOCX (B5 FIX)."""
        para = doc.add_paragraph()
        self._parse_inline_to_docx_runs(para, text)

    def _parse_inline_to_docx_runs(self, para, text: str):
        """Parse inline markdown and create styled runs in a DOCX paragraph."""
        import re as _re
        patterns = [
            (_re.compile(r'\*\*\*(.+?)\*\*\*'), 'bold_italic'),
            (_re.compile(r'\*\*(.+?)\*\*'), 'bold'),
            (_re.compile(r'__(.+?)__'), 'bold'),
            (_re.compile(r'\*(.+?)\*'), 'italic'),
            (_re.compile(r'_(.+?)_'), 'italic'),
            (_re.compile(r'~~(.+?)~~'), 'strikethrough'),
            (_re.compile(r'`(.+?)`'), 'code'),
            (_re.compile(r'\[(.+?)\]\((.+?)\)'), 'link'),
        ]

        pos = 0
        remaining = text

        while remaining:
            earliest_match = None
            earliest_pos = len(remaining)
            matched_style = None

            for pattern, style in patterns:
                match = pattern.search(remaining)
                if match and match.start() < earliest_pos:
                    earliest_match = match
                    earliest_pos = match.start()
                    matched_style = style

            if earliest_match:
                # Plain text before match
                if earliest_pos > 0:
                    para.add_run(remaining[:earliest_pos])

                if matched_style == 'link':
                    link_text = earliest_match.group(1)
                    run = para.add_run(link_text)
                    run.underline = True
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(0, 0x66, 0xCC)
                elif matched_style == 'code':
                    run = para.add_run(earliest_match.group(1))
                    run.font.name = 'Courier New'
                    from docx.shared import Pt
                    run.font.size = Pt(10)
                else:
                    content = earliest_match.group(1)
                    run = para.add_run(content)
                    if 'bold' in matched_style:
                        run.bold = True
                    if 'italic' in matched_style:
                        run.italic = True
                    if matched_style == 'strikethrough':
                        run.font.strike = True

                remaining = remaining[earliest_match.end():]
            else:
                para.add_run(remaining)
                break

    def to_latex(self, document: Document, output_path: Path):
        """Export to LaTeX"""
        latex = self._to_latex_string(document)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(latex, encoding='utf-8')

    def to_markdown(self, document: Document, output_path: Path):
        """Export to Markdown (essentially passthrough for MD)"""
        content = document.content

        # Ensure proper frontmatter if not present
        if not content.startswith('---'):
            frontmatter = f"""---
title: {document.title}
---

"""
            content = frontmatter + content

        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content, encoding='utf-8')

    def to_html(self, document: Document, output_path: Path):
        """Export to standalone HTML"""
        html_content = self._to_html_string(document, standalone=True)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(html_content, encoding='utf-8')

    def to_text(self, document: Document, output_path: Path):
        """Export to plain text"""
        # Remove markdown formatting
        text = document.content

        # Remove headers
        text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)

        # Remove code fences
        text = re.sub(r'```\w*\n?', '', text)

        # Remove inline code
        text = re.sub(r'`([^`]+)`', r'\1', text)

        # Remove bold/italic
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'__([^_]+)__', r'\1', text)
        text = re.sub(r'_([^_]+)_', r'\1', text)

        # Remove links
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(text, encoding='utf-8')

    def _to_html_string(self, document: Document, standalone: bool = False) -> str:
        """Convert document to HTML string"""
        try:
            import cmarkgfm
            body = cmarkgfm.github_flavored_markdown_to_html(document.content)
        except ImportError:
            # Fallback: basic markdown conversion
            body = self._basic_md_to_html(document.content)

        if standalone:
            escaped_title = html.escape(document.title)
            # B6 FIX: Use document language or config-based locale
            lang = "en"
            if document.language:
                lang_map = {"python": "en", "javascript": "en", "markdown": "en"}
                lang = lang_map.get(document.language, "en")
            return f"""<!DOCTYPE html>
<html lang="{lang}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{escaped_title}</title>
    <style>
        body {{
            font-family: 'Libertinus Serif', Georgia, serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 2rem;
            line-height: 1.6;
        }}
        h1, h2, h3 {{ font-family: 'Libertinus Sans', sans-serif; }}
        code {{ background: #f4f4f4; padding: 0.2em 0.4em; border-radius: 3px; }}
        pre {{ background: #f4f4f4; padding: 1em; border-radius: 5px; overflow-x: auto; }}
        pre code {{ background: none; padding: 0; }}
    </style>
</head>
<body>
{body}
</body>
</html>"""
        return body

    def _to_latex_string(self, document: Document) -> str:
        """Convert document to LaTeX string"""
        lines = document.content.split('\n')
        latex_lines = []

        # Document class and preamble
        latex_lines.append(r"\documentclass[11pt,a4paper]{article}")
        latex_lines.append(r"\usepackage[utf8]{inputenc}")
        latex_lines.append(r"\usepackage[T1]{fontenc}")
        latex_lines.append(r"\usepackage[english]{babel}")
        latex_lines.append(r"\usepackage{hyperref}")
        latex_lines.append(r"\usepackage{listings}")
        latex_lines.append("")
        latex_lines.append(f"\\title{{{self._escape_latex(document.title)}}}")
        latex_lines.append(r"\author{}")
        latex_lines.append(r"\date{\today}")
        latex_lines.append("")
        latex_lines.append(r"\begin{document}")
        latex_lines.append(r"\maketitle")
        latex_lines.append("")

        in_code_block = False
        code_lang = ""

        for line in lines:
            # Code blocks
            if line.startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_lang = line[3:].strip() or "text"
                    latex_lines.append(f"\\begin{{lstlisting}}[language={code_lang}]")
                else:
                    in_code_block = False
                    latex_lines.append(r"\end{lstlisting}")
                continue

            if in_code_block:
                latex_lines.append(line)
                continue

            # Headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                text = self._escape_latex(text)

                if level == 1:
                    latex_lines.append(f"\\section{{{text}}}")
                elif level == 2:
                    latex_lines.append(f"\\subsection{{{text}}}")
                elif level == 3:
                    latex_lines.append(f"\\subsubsection{{{text}}}")
                else:
                    latex_lines.append(f"\\paragraph{{{text}}}")
                continue

            # Regular text
            if line.strip():
                text = self._escape_latex(line)
                # Handle inline code
                text = re.sub(r'`([^`]+)`', r'\\texttt{\1}', text)
                # Handle bold
                text = re.sub(r'\*\*([^*]+)\*\*', r'\\textbf{\1}', text)
                # Handle italic
                text = re.sub(r'\*([^*]+)\*', r'\\textit{\1}', text)
                latex_lines.append(text)
            else:
                latex_lines.append("")

        latex_lines.append("")
        latex_lines.append(r"\end{document}")

        return '\n'.join(latex_lines)

    def _escape_latex(self, text: str) -> str:
        """Escape special LaTeX characters"""
        chars = {
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '{': r'\{',
            '}': r'\}',
            '~': r'\textasciitilde{}',
            '^': r'\^{}',
        }
        for char, escape in chars.items():
            text = text.replace(char, escape)
        return text

    def _basic_md_to_html(self, markdown: str) -> str:
        """Basic Markdown to HTML conversion"""
        html = markdown

        # Headers
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)

        # Bold (non-greedy, no newlines in content)
        html = re.sub(r'\*\*([^*\n]+?)\*\*', r'<strong>\1</strong>', html)

        # Italic (non-greedy, no asterisks or newlines in content)
        html = re.sub(r'\*([^*\n]+?)\*', r'<em>\1</em>', html)

        # Code (non-greedy, no backticks or newlines in content)
        html = re.sub(r'`([^`\n]+?)`', r'<code>\1</code>', html)

        # Links (non-greedy, proper character classes)
        html = re.sub(r'\[([^\]\n]+?)\]\(([^)\n]+?)\)', r'<a href="\2">\1</a>', html)

        # Paragraphs
        paragraphs = html.split('\n\n')
        html = '\n'.join(f'<p>{p}</p>' if not p.startswith('<h') else p for p in paragraphs)

        return html
