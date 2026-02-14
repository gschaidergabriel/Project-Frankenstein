"""
DOCX Parser for document ingestion
"""

import logging
from pathlib import Path
from typing import List, Dict

# Module-level docx import with error handling (DRY)
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


def _check_docx_available():
    """Check if python-docx is available and raise if not"""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")


def _parse_heading_level(style_name: str) -> int:
    """Parse heading level from style name, handling non-numeric gracefully"""
    try:
        level_str = style_name.replace('Heading ', '').strip()
        if level_str and level_str.isdigit():
            return int(level_str)
        return 1  # Default to level 1 for non-numeric headings
    except (ValueError, AttributeError):
        return 1


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text from DOCX file"""
    _check_docx_available()

    docx_path = Path(docx_path)

    # File existence check before Document()
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    try:
        doc = Document(str(docx_path))
    except Exception as e:
        logger.error(f"Failed to open DOCX file {docx_path}: {type(e).__name__}: {e}")
        raise

    paragraphs = []

    try:
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Check if para.style is None before accessing .name
                if para.style is not None and para.style.name and para.style.name.startswith('Heading'):
                    level = _parse_heading_level(para.style.name)
                    prefix = '#' * level
                    paragraphs.append(f"{prefix} {text}")
                else:
                    paragraphs.append(text)
    except Exception as e:
        logger.error(f"Failed to parse paragraphs from {docx_path}: {type(e).__name__}: {e}")
        raise

    return "\n\n".join(paragraphs)


def extract_structure_from_docx(docx_path: Path) -> List[Dict]:
    """Extract document structure (headings, sections)"""
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not available, cannot extract document structure")
        return []

    docx_path = Path(docx_path)

    # File existence check before Document()
    if not docx_path.exists():
        logger.error(f"DOCX not found for structure extraction: {docx_path}")
        return []

    try:
        doc = Document(str(docx_path))
        structure = []

        for i, para in enumerate(doc.paragraphs):
            # Add None check for para.style
            if para.style is not None and para.style.name and para.style.name.startswith('Heading'):
                level = _parse_heading_level(para.style.name)
                structure.append({
                    'level': level,
                    'title': para.text.strip(),
                    'index': i
                })

        return structure
    except Exception as e:
        logger.error(f"Failed to extract structure from {docx_path}: {type(e).__name__}: {e}")
        return []


def get_docx_metadata(docx_path: Path) -> dict:
    """Extract DOCX metadata"""
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not available, cannot extract DOCX metadata")
        return {}

    docx_path = Path(docx_path)

    # File existence check before Document()
    if not docx_path.exists():
        logger.error(f"DOCX not found for metadata extraction: {docx_path}")
        return {}

    try:
        doc = Document(str(docx_path))
        props = doc.core_properties

        return {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
        }
    except Exception as e:
        logger.error(f"Failed to extract DOCX metadata from {docx_path}: {type(e).__name__}: {e}")
        return {}
